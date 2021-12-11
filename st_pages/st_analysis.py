import numpy as np
import pandas as pd
import streamlit as st
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.stem import SnowballStemmer
from nltk.tokenize import RegexpTokenizer
from nltk.corpus import stopwords
from spacy import displacy
nltk.download('stopwords')
from docx import Document
import spacy

from st_pages.wordcloud import st_wordcloud
from st_pages.st_utils import st_img

nlp = spacy.load('ru_core_news_md')
stop_words = stopwords.words('russian')
stem = SnowballStemmer('russian')
punc = RegexpTokenizer(r'\w+')


def tables_2d_frame(data):
    res = []
    for table in data.tables:
        out_table = [[0 for _ in range(len(table.columns))] for _ in range(len(table.rows))]
        for index_r, row in enumerate(table.rows):
            for index_c, cell in enumerate(row.cells):
                if cell.text:
                    out_table[index_r][index_c] = cell.text
        res.append(pd.DataFrame(out_table))
    return res


def preproc_raw(raw, lemmatize=True):
    curr_raw = word_tokenize(raw, 'russian')
    if not lemmatize:
        return curr_raw
    new_raw = " ".join([stem.stem(i) for i in curr_raw if i not in stop_words and i.isalpha()])
    return new_raw


def parag_2d_text(para, lemmatize=True):
    para = np.array([p.text.replace("\t", " ").strip() for p in para.paragraphs])
    res = []
    for i in para:
        if i != '' and i != ' ':
            res.append(preproc_raw(i.lower(), lemmatize=lemmatize))
    return np.array(res)


def st_load_docx():
    from st_pages.spacy_formatter import format_string_as_spacy
    doc_file = st.file_uploader("Загрузите документ для анализа", type=["docx"])

    if doc_file is not None:
        content = doc_file.read()
        with open("tmp.docx", "wb+") as writer:
            writer.write(content)
        doc = Document("tmp.docx")
        st.write(type(doc))
        table, text = tables_2d_frame(doc), parag_2d_text(doc, lemmatize=False)
        requirements_txt = []
        requirements_raw = []

        for para in doc.paragraphs:
            format_string_as_spacy(para.text)
            if "требован" in para.text:
                raw_txt = [t for t in para.text.split("Преимущества участия")[0].split("\n") if t != ""]
                cleaned_txt = preproc_raw(para.text.replace("\t", " ").strip().lower(),
                                          lemmatize=False)
                if not cleaned_txt == "":
                    requirements_txt.append(cleaned_txt)
                    requirements_raw.append(raw_txt)

        st.sidebar.text("Облако тегов")
        st_wordcloud(requirements_txt, stop_words)
        st_img("wc.png", sidebar=True, width=300)

        st.sidebar.text("Требования")
        st.sidebar.write(requirements_raw)

    # test_doc = "test.docx"
    # document = Document(test_doc)
    # table, text = tables_2d_frame(document), parag_2d_text(document, lemmatize=False)
